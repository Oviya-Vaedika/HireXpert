import streamlit as st
import google.generativeai as genai
import pypdf 
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from datetime import datetime
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from io import BytesIO
from fpdf import FPDF

# Load Gemini API key from Streamlit Secrets
genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

# Load the model instance
model = genai.GenerativeModel("gemini-2.5-flash")

# App Page Layout Setup
st.set_page_config(page_title="HireXpert", page_icon="🤖", layout="wide")

# Initialize Session States safely across 6 Student Application Modules
if "page" not in st.session_state:
    st.session_state.page = "home"
if "jobs" not in st.session_state:
    st.session_state.jobs = []
if "education" not in st.session_state:
    st.session_state.education = []
if "projects" not in st.session_state:
    st.session_state.projects = []
if "personal_info" not in st.session_state:
    st.session_state.personal_info = {"name": "", "email": "", "phone": "", "linkedin": ""}
if "skills_info" not in st.session_state:
    st.session_state.skills_info = {"skills": ""}
if "template" not in st.session_state:
    st.session_state.template = "Classic Corporate"

# Helper Function to completely strip out markdown formatting asterisks
def strip_markdown_stars(text):
    if not text:
        return ""
    # Strip markdown bold markers (**bold** -> bold)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    # Strip plain asterisks and hanging bullet points cleanly
    text = re.sub(r"\* ?", "", text)
    return text.strip()

# Extraction Helper Function for Analyzer Page
def extract_text(uploaded_file):
    text = ""
    try:
        if uploaded_file.type == "application/pdf":
            pdf_reader = pypdf.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs if para.text])
        elif uploaded_file.type == "text/plain":
            uploaded_file.seek(0) 
            text = uploaded_file.read().decode("utf-8")
        else:
            return None
    except Exception as e:
        print(f"Extraction error: {e}") 
        return None
    return text.strip()
# 📊 SMART ATS SCORING: Falls back to AI evaluation if the JD is too short
def calculate_score(resume, job_desc):
    if not resume.strip() or not job_desc.strip():
        return 0.0
        
    if len(job_desc.split()) < 10:
        try:
            score_prompt = f"""
            You are an ATS algorithm. Compare this resume against the target role title.
            Target Role: {job_desc}
            Resume Content: {resume}
            Output ONLY a single number representing the percentage alignment (e.g., 78.5). No words.
            """
            ai_score_response = model.generate_content(score_prompt).text.strip()
            cleaned_score = re.findall(r"\d+\.\d+|\d+", ai_score_response)
            if cleaned_score:
                return round(float(cleaned_score[0]), 2)
        except:
            pass 

    try:
        vectorizer = TfidfVectorizer(token_pattern=r'\b\w+\b')
        tfidf_matrix = vectorizer.fit_transform([resume, job_desc])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
        return round(similarity * 100, 2)
    except:
        return 0.0
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

# Pure Python Layout Engine to render PDF streams instantly from memory
def generate_pdf_file():
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    p_info = st.session_state.personal_info
    s_info = st.session_state.skills_info
    template = st.session_state.template
    
    if template == "Classic Corporate":
        pdf.set_font("Times", size=11)
        primary_color = (0, 0, 0)
        secondary_color = (90, 90, 90)
    elif template == "Modern Minimalist":
        pdf.set_font("Arial", size=10)
        primary_color = (44, 62, 80)
        secondary_color = (127, 140, 141)
    elif template == "Tech Startup":
        pdf.set_font("Courier", size=11)
        primary_color = (26, 188, 156)
        secondary_color = (52, 73, 94)
    else: 
        pdf.set_font("Times", size=11)
        primary_color = (139, 0, 0)
        secondary_color = (60, 60, 60)
        
    # Header Module
    pdf.set_text_color(*primary_color)
    pdf.set_font(pdf.current_font.family, 'B', 22)
    pdf.cell(0, 12, strip_markdown_stars(p_info['name'] if p_info['name'] else "Resume Profile"), ln=True, align='C' if template == "Classic Corporate" else 'L')
    
    pdf.set_text_color(*secondary_color)
    pdf.set_font(pdf.current_font.family, size=10)
    pdf.cell(0, 6, f"Email: {p_info['email']} | Phone: {p_info['phone']} | LinkedIn: {p_info['linkedin']}", ln=True, align='C' if template == "Classic Corporate" else 'L')
    pdf.ln(5)
    
    # Skills Module
    if s_info['skills']:
        pdf.set_text_color(*primary_color)
        pdf.set_font(pdf.current_font.family, 'B', 14)
        pdf.cell(0, 10, "Core Competencies & Skills" if template != "Tech Startup" else "// Core Tech Stack", ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font(pdf.current_font.family, size=11)
        pdf.multi_cell(0, 6, strip_markdown_stars(s_info['skills']))
        pdf.ln(4)
        
    # Professional Record Module
    if st.session_state.jobs:
        pdf.set_text_color(*primary_color)
        pdf.set_font(pdf.current_font.family, 'B', 14)
        pdf.cell(0, 10, "Professional Experience" if template != "Tech Startup" else "// Technical Deployments", ln=True)
        
        for job in st.session_state.jobs:
            pdf.set_text_color(*secondary_color)
            pdf.set_font(pdf.current_font.family, 'B', 12)
            pdf.cell(0, 8, f"{strip_markdown_stars(job['title'])} - {strip_markdown_stars(job['company'])}", ln=True)
            pdf.set_font(pdf.current_font.family, 'I', 10)
            pdf.cell(0, 6, f"{job['dates']} | {job['location']}", ln=True)
            
            pdf.set_text_color(0, 0, 0)
            pdf.set_font(pdf.current_font.family, size=11)
            for bullet in job['bullets'].split('\n'):
                if bullet.strip():
                    cleaned_b = strip_markdown_stars(bullet)
                    if cleaned_b:
                        pdf.multi_cell(0, 6, f"  - {cleaned_b}")
            pdf.ln(2)
            
    # Projects Module
    if st.session_state.projects:
        pdf.set_text_color(*primary_color)
        pdf.set_font(pdf.current_font.family, 'B', 14)
        pdf.cell(0, 10, "Projects & Internships" if template != "Tech Startup" else "// Practical Implementations", ln=True)
        
        for proj in st.session_state.projects:
            pdf.set_text_color(0, 0, 0)
            pdf.set_font(pdf.current_font.family, 'B', 11)
            pdf.cell(0, 7, f"{strip_markdown_stars(proj['title'])} ({proj['timeline']})", ln=True)
            pdf.set_font(pdf.current_font.family, size=11)
            pdf.multi_cell(0, 6, strip_markdown_stars(proj['details']))
            pdf.ln(2)
            
    # Education Module
    if st.session_state.education:
        pdf.set_text_color(*primary_color)
        pdf.set_font(pdf.current_font.family, 'B', 14)
        pdf.cell(0, 10, "Education Qualifications" if template != "Tech Startup" else "// Academic Framework", ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font(pdf.current_font.family, size=11)
        for edu in st.session_state.education:
            pdf.multi_cell(0, 6, f"Education: {strip_markdown_stars(edu['degree'])} from {strip_markdown_stars(edu['school'])} ({edu['date']})")
            if edu['details']:
                pdf.multi_cell(0, 5, f"Details: {strip_markdown_stars(edu['details'])}")
            pdf.ln(2)

    return bytes(pdf.output())

# Cleaned Word File Export Engine
def generate_docx_file():
    doc = Document()
    p_info = st.session_state.personal_info
    s_info = st.session_state.skills_info
    template = st.session_state.template
    
    style = doc.styles['Normal']
    style.font.name = 'Arial' if template == "Modern Minimalist" else 'Times New Roman'
    style.font.size = Pt(11)
    
    np = doc.add_paragraph()
    n_run = np.add_run(strip_markdown_stars(p_info['name']))
    n_run.font.size = Pt(22)
    n_run.bold = True
    
    doc.add_paragraph(f"Email: {p_info['email']} | Phone: {p_info['phone']} | LinkedIn: {p_info['linkedin']}")
    
    if s_info['skills']:
        doc.add_heading('Core Competencies & Skills', level=1)
        doc.add_paragraph(strip_markdown_stars(s_info['skills']))
        
    if st.session_state.jobs:
        doc.add_heading('Professional Experience', level=1)
        for job in st.session_state.jobs:
            p = doc.add_paragraph()
            p.add_run(f"{strip_markdown_stars(job['title'])} — {strip_markdown_stars(job['company'])} ({job['dates']})\n").bold = True
            for b in job['bullets'].split('\n'):
                if b.strip():
                    cleaned_bullet = strip_markdown_stars(b)
                    if cleaned_bullet:
                        doc.add_paragraph(style='List Bullet').add_run(cleaned_bullet)
                    
    if st.session_state.projects:
        doc.add_heading('Projects & Internships', level=1)
        for proj in st.session_state.projects:
            p = doc.add_paragraph()
            p.add_run(f"{strip_markdown_stars(proj['title'])} ({proj['timeline']})\n").bold = True
            p.add_run(strip_markdown_stars(proj['details']))

    if st.session_state.education:
        doc.add_heading('Education Qualifications', level=1)
        for edu in st.session_state.education:
            doc.add_paragraph(f"• {strip_markdown_stars(edu['degree'])} from {strip_markdown_stars(edu['school'])} ({edu['date']})")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
# SCREEN INTERACTION GATEWAY
if st.session_state.page == "home":
    st.title("🤖 Welcome to HireXpert")
    st.write("What would you like to do today?")
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("📊 Analyze Existing Resume")
        if st.button("🚀 Go to Analyzer", use_container_width=True):
            st.session_state.page = "analyzer"
            st.rerun()
    with col2:
        st.subheader("📄 Build a New Resume")
        if st.button("🛠️ Go to Builder", use_container_width=True):
            st.session_state.page = "builder"
            st.rerun()

elif st.session_state.page == "analyzer":
    if st.button("⬅️ Back to Home Menu"):
        st.session_state.page = "home"
        st.rerun()
    st.title("🤖 HireXpert - Resume Analyzer")
    resume_file = st.file_uploader("Upload Resume File", type=["pdf", "docx", "txt"])
    job_desc_input = st.text_area("Paste Target Profile Criteria / Job Description")
    if st.button("Run Profile Analysis"):
        if resume_file and job_desc_input.strip():
            resume_text = extract_text(resume_file)
            score = calculate_score(resume_text, job_desc_input)
            st.subheader(f"📊 Semantic Fit Score: {score}%")
            st.progress(score / 100.0)
            with st.spinner("Analyzing data layers..."):
                st.write(model.generate_content(f"Analyze: {resume_text} Against Criteria: {job_desc_input}").text)

elif st.session_state.page == "builder":
    if st.button("⬅️ Back to Home Menu"):
        st.session_state.page = "home"
        st.rerun()
        
    st.title("📄 Multi-History AI Resume Builder")
    col1, col2 = st.columns(2)
    
    with col1:
        # MODULE 1: PERSONAL DETAILS
        st.subheader("👤 Box 1: Personal Details")
        with st.form("personal_form"):
            f_name = st.text_input("Full Name", value=st.session_state.personal_info["name"])
            e_mail = st.text_input("Email Address", value=st.session_state.personal_info["email"])
            p_hone = st.text_input("Phone Number", value=st.session_state.personal_info["phone"])
            l_inkedin = st.text_input("LinkedIn Profile URL", value=st.session_state.personal_info["linkedin"])
            if st.form_submit_button("💾 Save Personal Data"):
                st.session_state.personal_info = {"name": f_name, "email": e_mail, "phone": p_hone, "linkedin": l_inkedin}
                st.success("Personal Details saved!")

        # MODULE 2: WORK EXPERIENCE WITH APPEND
        st.subheader("💼 Box 2: Professional Experience")
        with st.form("experience_form", clear_on_submit=True):
            job_title = st.text_input("Job Title / Position")
            company = st.text_input("Company / Organization")
            dates = st.text_input("Employment Dates / Timeline")
            location = st.text_input("Location")
            raw_desc = st.text_area("Responsibilities Notes (AI Optimizes this)")
            if st.form_submit_button("➕ Add Work Position Profile"):
                if job_title and company:
                    with st.spinner("Optimizing job points using AI metrics..."):
                        bullets = model.generate_content(f"Convert into 3 powerful resume bullet points starting with standard (-). Strip out any wrapping asterisks or markdown decorators. Text: {raw_desc}").text.strip() if raw_desc.strip() else "- Performance managed roles."
                    st.session_state.jobs.append({"title": job_title, "company": company, "dates": dates, "location": location, "bullets": bullets})
                    st.rerun()

        # MODULE 3: EDUCATION
        st.subheader("🎓 Box 3: Academic Qualifications")
        with st.form("education_form", clear_on_submit=True):
            school_name = st.text_input("Institution / University Name")
            degree_name = st.text_input("Degree / Major")
            grad_date = st.text_input("Graduation Timeline")
            edu_details = st.text_input("Additional Info (GPA, Awards)")
            if st.form_submit_button("➕ Add Academic History Entry"):
                if school_name and degree_name:
                    st.session_state.education.append({"school": school_name, "degree": degree_name, "date": grad_date, "details": edu_details})
                    st.rerun()

        # MODULE 4: PROJECTS OR INTERNSHIPS
        st.subheader("🚀 Box 4: Projects & Internships")
        with st.form("project_form", clear_on_submit=True):
            p_title = st.text_input("Project / Internship Title")
            p_time = st.text_input("Timeline / Duration")
            p_desc = st.text_area("Scope & Technologies Used Summary")
            if st.form_submit_button("➕ Add Project Portfolio Entry"):
                if p_title:
                    st.session_state.projects.append({"title": p_title, "timeline": p_time, "details": p_desc})
                    st.rerun()

        # MODULE 5: SKILLS WITH INTELLIGENT ACCELERATOR/EXPANDER
        st.subheader("🛠️ Box 5: Core Competencies & Skills Optimization")
        with st.form("skills_form"):
            skills_input = st.text_area("Type raw paragraph notes or technologies list here:")
            if st.form_submit_button("🤖 Abbreviate & Expand via Semantic Engine"):
                if skills_input.strip():
                    with st.spinner("Analyzing skill matrices..."):
                        skill_prompt = f"Extract the technical competencies from this text. Convert into a neat comma-separated list of keywords. Strip out any markdown formatting asterisks. Text: {skills_input}"
                        optimized_skills = model.generate_content(skill_prompt).text.strip()
                        st.session_state.skills_info = {"skills": optimized_skills}
                        st.success("Skills matrix optimized!")
                else:
                    st.error("Please enter a few keywords or paragraph blocks.")

    # MODULE 6: LIVE LAYOUT ENGINE SELECTION DISPLAY PREVIEW
    with col2:
        st.subheader("🎨 Box 6: Template Strategy Framework Layout Selector")
        st.session_state.template = st.selectbox(
            "Select Layout Style Blueprint Structure:", 
            ["Classic Corporate", "Modern Minimalist", "Tech Startup", "Executive Elegance"]
        )
        
        st.markdown("---")
        st.subheader("👁️ Live Structural Mockup Viewport")
        
        clean_name = strip_markdown_stars(st.session_state.personal_info['name'])
        
        if st.session_state.template == "Modern Minimalist":
            m_col1, m_col2 = st.columns(2)
            with m_col1:
                st.markdown(f"### **{clean_name}**")
                st.caption(f"📱 {st.session_state.personal_info['phone']}\n\n✉️ {st.session_state.personal_info['email']}")
                if st.session_state.skills_info['skills']:
                    st.markdown("**EXPERTISE**")
                    for s in st.session_state.skills_info['skills'].split(','):
                        st.markdown(f"- {strip_markdown_stars(s)}")
            with m_col2:
                if st.session_state.jobs:
                    st.markdown("#### **PROFESSIONAL REFERENCES**")
                    for j in st.session_state.jobs:
                        st.markdown(f"**{strip_markdown_stars(j['title'])}** at *{strip_markdown_stars(j['company'])}* ({j['dates']})")
                        st.write(strip_markdown_stars(j['bullets']))
                        
        elif st.session_state.template == "Tech Startup":
            st.success(f"📌 TOP HEADER ACTIVE // {clean_name.upper()}")
            st.write(f"⚙️ Info Matrix: {st.session_state.personal_info['email']} | {st.session_state.personal_info['phone']}")
            st.markdown("---")
            if st.session_state.skills_info['skills']:
                st.write(f"**// Core Stack:** {strip_markdown_stars(st.session_state.skills_info['skills'])}")
            for j in st.session_state.jobs:
                st.write(f"💼 **// Deployment:** {strip_markdown_stars(j['title'])} @ {strip_markdown_stars(j['company'])}")
                st.write(strip_markdown_stars(j['bullets']))

        else:
            align_prefix = "###" if st.session_state.template == "Classic Corporate" else "### <div style='text-align: right;'>"
            st.markdown(f"{align_prefix} **{clean_name}** </div>", unsafe_allow_html=True)
            st.markdown(f"<div style='text-align: center;'> {st.session_state.personal_info['email']} | {st.session_state.personal_info['phone']} </div>", unsafe_allow_html=True)
            st.markdown("---")
            if st.session_state.skills_info['skills']:
                st.write(f"**CORE COMPETENCIES:** {strip_markdown_stars(st.session_state.skills_info['skills'])}")
            for j in st.session_state.jobs:
                st.write(f"**{strip_markdown_stars(j['title'])}** — {strip_markdown_stars(j['company'])} ({j['dates']})")
                st.write(strip_markdown_stars(j['bullets']))
                
        if st.session_state.projects:
            st.markdown("---")
            st.markdown("#### **Project Subsystem Frameworks**")
            for p in st.session_state.projects:
                st.write(f"🚀 **{strip_markdown_stars(p['title'])}** ({p['timeline']}): {strip_markdown_stars(p['details'])}")
                
        if st.session_state.education:
            st.markdown("---")
            st.markdown("#### **Academic Credentials**")
            for e in st.session_state.education:
                st.write(f"🎓 {strip_markdown_stars(e['degree'])} from {strip_markdown_stars(e['school'])}")

              st.markdown("---")
        if st.session_state.personal_info['name']:
            btn_col1, btn_col2 = st.columns(2)
            with btn_col1:
                st.download_button(
                    label=f"📥 Download Direct PDF Asset",
                    data=generate_pdf_file(),
                    file_name=f"HireXpert_{st.session_state.template.replace(' ', '_')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
                
            with btn_col2:
                st.download_button(
                    label=f"📄 Download Word File (.docx)",
                    data=generate_docx_file(),
                    file_name=f"HireXpert_{st.session_state.template.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        else:
            st.warning("Complete Box 1 (Personal Details) with a validated name to unlock PDF & Word download engines.")

        # --- STUDENT COMPONENT APP FOOTER ---
        st.markdown("---")
        st.markdown(
            """
            <div style='text-align: center; color: #7f8c8d; font-size: 14px; padding-top: 10px;'>
                🤖 Made with ❤️ by <b>HireXpert Student Team</b> | © 2026 AI Resume Engine
            </div>
            """, 
            unsafe_allow_html=True
        )

